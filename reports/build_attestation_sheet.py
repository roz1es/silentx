from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path("/Users/roz1es/Downloads/Аттестационный лист ПП 04.01 ВТ-Х-ХХ.docx")
OUT = Path("/Users/roz1es/Downloads/Аттестационный лист ПП 04.01 БренксЧат готовый.docx")
FONT = "Times New Roman"


def set_run_font(run, size: float = 12, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, size: float = 12, bold: bool = False, align=None) -> None:
    cell.text = ""
    lines = text.split("\n")
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    for idx, line in enumerate(lines):
        if idx:
            para.add_run().add_break()
        run = para.add_run(line)
        set_run_font(run, size, bold)


def remove_row(table, row_index: int) -> None:
    table._tbl.remove(table.rows[row_index]._tr)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def clone_last_row(table):
    return deepcopy(table.rows[-1]._tr)


def main() -> None:
    doc = Document(SRC)

    for style in doc.styles:
        if hasattr(style, "font"):
            style.font.name = FONT
            style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # ФИО и группа оставлены как поля для ручного заполнения.
    if len(doc.tables) > 1:
        set_cell_text(doc.tables[1].cell(0, 0), "__________________________________________", 12)

    if len(doc.tables) > 2:
        t = doc.tables[2]
        set_cell_text(t.cell(0, 1), "____", 12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t.cell(1, 4), "ВТ-Х-ХХ", 12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Виды работ по фактическому проекту BrenksChat.
    works = [
        (
            "Анализ требований к веб-мессенджеру «БренксЧат», определение пользовательских сценариев, "
            "проектирование клиент-серверной архитектуры приложения.",
            "8",
        ),
        (
            "Проектирование структуры данных и MySQL-совместимой базы данных: пользователи, чаты, участники, "
            "сообщения, сессии, push-подписки и E2EE-устройства.",
            "10",
        ),
        (
            "Разработка серверной части на Node.js, Express и Socket.IO: REST API, realtime-события, обработка "
            "сообщений, прав доступа и состояний пользователей.",
            "14",
        ),
        (
            "Разработка клиентской части на React, TypeScript, Vite и Tailwind CSS: страницы входа, мессенджера, "
            "админ-панели, список чатов, окно переписки, профили и модальные окна.",
            "14",
        ),
        (
            "Реализация авторизации и аутентификации: регистрация, вход, подтверждение почты кодом, сброс пароля, "
            "JWT/cookie-сессии, Argon2id-хэширование паролей и роль администратора.",
            "8",
        ),
        (
            "Реализация обмена сообщениями и медиа: личные чаты, группы, каналы, ответы, реакции, редактирование, "
            "удаление, пересылка, изображения, файлы, голосовые сообщения и видеокружки.",
            "8",
        ),
        (
            "Настройка постоянного хранения данных, push-уведомлений, WebRTC-звонков и серверной инфраструктуры: "
            "Nginx reverse proxy, systemd-сервис, переменные окружения и резервное копирование.",
            "6",
        ),
        (
            "Тестирование и отладка проекта: запуск серверных автоматизированных тестов, production-сборка, "
            "проверка TypeScript, исправление ошибок и подготовка проектной документации.",
            "4",
        ),
    ]

    work_table = doc.tables[5]
    header = work_table.rows[0]
    total_template = clone_last_row(work_table)
    while len(work_table.rows) > 1:
        remove_row(work_table, 1)
    for work, hours in works:
        row = work_table.add_row()
        prevent_row_split(row)
        set_cell_text(row.cells[0], work, 10.2)
        set_cell_text(row.cells[1], hours, 10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    work_table._tbl.append(total_template)
    for row in work_table.rows:
        prevent_row_split(row)
    total_row = work_table.rows[-1]
    set_cell_text(total_row.cells[0], "Итого часов:", 12, True, WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(total_row.cells[1], "72", 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header.cells[0], "Виды работ", 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header.cells[1], "Объем выполненных работ (часов)", 12, True, WD_ALIGN_PARAGRAPH.CENTER)

    if len(doc.tables) > 7:
        set_cell_text(
            doc.tables[7].cell(0, 0),
            "3. База прохождения производственной практики\n"
            "Учебный проект по разработке веб-мессенджера «БренксЧат» на базе собственного рабочего места "
            "студента и серверной среды VPS (Node.js, React, MySQL, Nginx).",
            12,
        )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
