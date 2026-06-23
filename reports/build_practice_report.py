from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path("/Users/roz1es/Downloads/messengercursor 2/messengercursor")
OUT = Path("/Users/roz1es/Downloads/Отчет ПП 04.01 БренксЧат готовый.docx")

FONT = "Times New Roman"
PAGE_WIDTH_CM = 17.5


def set_style_font(style, size: float, bold: bool = False, color: str | None = None) -> None:
    font = style.font
    font.name = FONT
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic:
        run.italic = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    set_run_font(run, 12)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(10)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in [("Heading 1", 16), ("Heading 2", 16), ("Heading 3", 15)]:
        style = styles[name]
        set_style_font(style, size, True)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        set_style_font(style, 14)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)


def p(doc: Document, text: str = "", style: str | None = None, align=None, first_indent=True):
    para = doc.add_paragraph(style=style)
    if text:
        run = para.add_run(text)
        set_run_font(run, 14)
    if align is not None:
        para.alignment = align
    if not first_indent:
        para.paragraph_format.first_line_indent = Cm(0)
    return para


def h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    set_run_font(run, 14)
    return para


def numbered(doc: Document, text: str):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    set_run_font(run, 14)
    return para


def manual_numbered(doc: Document, number: int, text: str, left: bool = False):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0.75)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(f"{number}. {text}")
    set_run_font(run, 14)
    return para


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    cell.width = Cm(cm)


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) / 2.54 * 1440)))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def write_cell(cell, text: str, bold: bool = False, size: float = 11) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = 1.15
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, size, bold)


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float]):
    cap = p(doc, caption, first_indent=False)
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(4)
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    set_table_width(tbl, widths)
    set_cell_margins(tbl)
    for idx, header in enumerate(headers):
        shade_cell(tbl.rows[0].cells[idx], "E8EEF5")
        write_cell(tbl.rows[0].cells[idx], header, True, 11)
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            write_cell(cells[idx], value, False, 10.5 if len(value) > 120 else 11)
    set_table_width(tbl, widths)
    return tbl


def screenshot_marker(doc: Document, text: str) -> None:
    para = p(doc, f"[Вставить скриншот: {text}]", first_indent=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        set_run_font(run, 13, italic=True)
        run.font.color.rgb = RGBColor(90, 90, 90)


def title_page(doc: Document) -> None:
    lines = [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "федеральное государственное бюджетное образовательное учреждение высшего образования",
        "«Российский экономический университет имени Г. В. Плеханова»",
        "МОСКОВСКИЙ ПРИБОРОСТРОИТЕЛЬНЫЙ ТЕХНИКУМ",
    ]
    for line in lines:
        para = p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
        for run in para.runs:
            set_run_font(run, 12, True if line.startswith("МОСКОВСКИЙ") else False)

    doc.add_paragraph()
    doc.add_paragraph()
    para = p(doc, "ОТЧЕТ", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for run in para.runs:
        set_run_font(run, 18, True)
    para = p(doc, "по производственной практике", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    para = p(doc, "ПП.04.01 Производственная практика", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    para = p(doc, "по специальности 09.02.09 «Веб-разработка»", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    para = p(doc, "Тема: разработка веб-мессенджера «БренксЧат»", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for paragraph in doc.paragraphs[-4:]:
        for run in paragraph.runs:
            set_run_font(run, 14, paragraph.text.startswith("Тема"))

    doc.add_paragraph()
    doc.add_paragraph()
    info = [
        "Студент: ________________________________",
        "Группа: ВТ-Х-ХХ",
        "Руководитель практической подготовки от профильной организации: ________________________________",
        "Руководитель практической подготовки от техникума: ________________________________",
    ]
    for line in info:
        para = p(doc, line, first_indent=False)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for _ in range(5):
        doc.add_paragraph()
    para = p(doc, "Москва, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    doc.add_page_break()


def contents(doc: Document) -> None:
    h(doc, "Содержание", 1)
    items = [
        "Результаты анализа проекта",
        "Введение",
        "1 Цели и задачи практики",
        "2 База практики",
        "3 Содержание практики",
        "4 Охрана труда",
        "Вывод",
        "Список использованных источников",
    ]
    for item in items:
        p(doc, item, first_indent=False)
    doc.add_page_break()


def main() -> None:
    doc = Document()
    configure_document(doc)
    title_page(doc)
    contents(doc)

    h(doc, "Результаты анализа проекта", 1)
    p(
        doc,
        "В ходе анализа проекта установлено, что разработанное приложение представляет собой веб-мессенджер "
        "«БренксЧат» с клиент-серверной архитектурой. Клиентская часть реализована как одностраничное "
        "приложение на React и TypeScript, серверная часть — на Node.js, Express и Socket.IO. Для постоянного "
        "хранения данных используется MySQL-совместимая база данных через библиотеку mysql2; при отсутствии "
        "MySQL предусмотрен файловый вариант хранения состояния в JSON-файле.",
    )
    bullet(doc, "Назначение проекта: обеспечение обмена сообщениями в реальном времени, работы личных чатов, групп, каналов, медиа-сообщений, голосовых и видеозвонков, а также администрирования пользователей.")
    bullet(doc, "Клиентские технологии: React 18, TypeScript, Vite, Tailwind CSS, React Router, Socket.IO Client, Web APIs браузера, Service Worker и PWA manifest.")
    bullet(doc, "Серверные технологии: Node.js, Express, Socket.IO, MySQL2, JSON Web Token, Argon2id, Helmet, CORS, express-rate-limit, Nodemailer, Web Push, UUID.")
    bullet(doc, "Криптографические и защитные механизмы: хэширование паролей Argon2id, cookie-сессии с JWT, ограничение частоты авторизационных запросов, подтверждение почты кодом, резервное копирование E2EE-ключей, шифрование текстов личных сообщений через libsodium.")
    bullet(doc, "Архитектура: SPA-клиент, REST API, WebSocket-слой для realtime-событий, слой хранения состояния, MySQL-схема, Nginx reverse proxy с HTTPS и отдельные deploy-конфигурации.")
    bullet(doc, "Основные клиентские модули: AuthContext, MessengerContext, CallContext, ThemeContext, ChatWallpaperContext; компоненты ChatList, ChatWindow, MessageBubble, MessageInput, NewChatModal, ProfileModal, PhotoViewer, VoiceVideoMedia.")
    bullet(doc, "Основные серверные модули: index.ts, socketHandlers.ts, store.ts, persist.ts, auth.ts, sessions.ts, password.ts, emailAuth.ts, email.ts, e2eeDevices.ts, pushNotifications.ts.")
    bullet(doc, "Структура базы данных: users, chats, chat_participants, messages, user_chat_muted, user_chat_pinned, push_subscriptions, auth_sessions, user_e2ee_devices, user_e2ee_key_backups, app_state.")
    bullet(doc, "Тестирование: в серверной части обнаружены автоматизированные тесты для Argon2id-паролей, сессий, E2EE-устройств и миграции учетных данных; команда npm test -w server выполнена успешно, 8 тестов пройдено.")
    bullet(doc, "Сборка: команда npm run build выполнена успешно; клиент собран Vite, сервер скомпилирован TypeScript-компилятором.")

    missing_rows = [
        ["ФИО студента, номер курса, точная группа", "В шаблонах оставлены поля-заполнители.", "Вписать фактические данные студента и группы перед сдачей отчета."],
        ["Название профильной организации", "В проекте и шаблонах не указано.", "Использовать формулировку о прохождении практики в рамках учебного проекта либо вписать организацию, если она была назначена техникумом."],
        ["Точные характеристики компьютера", "В репозитории отсутствует паспорт оборудования.", "Заполнить по фактическому ПК: процессор, ОЗУ, накопитель, монитор, периферия."],
        ["Точная версия MySQL на сервере", "В коде используется mysql2 и MySQL-совместимый синтаксис, но версия СУБД не зафиксирована.", "Указать версию из phpMyAdmin или команды SELECT VERSION()."],
        ["ФИО руководителей практики", "В предоставленных документах поля пустые.", "Заполнить по данным образовательной организации и базы практики."],
    ]
    table(
        doc,
        "Таблица 1 — Данные, отсутствующие в проектных материалах",
        ["Данные", "Что найдено", "Комментарий для заполнения"],
        missing_rows,
        [4.4, 5.8, 7.3],
    )

    h(doc, "Введение", 1)
    p(
        doc,
        "Производственная практика ПП.04.01 по специальности 09.02.09 «Веб-разработка» была направлена "
        "на закрепление профессиональных навыков разработки веб-приложений и получение практического "
        "опыта полного цикла создания программного продукта. В качестве практического задания был разработан "
        "мессенджер «БренксЧат», предназначенный для обмена сообщениями, управления чатами, отправки медиафайлов, "
        "организации голосовой и видеосвязи, а также администрирования пользователей.",
    )
    p(
        doc,
        "Сведения о конкретной профильной организации в исходных документах и проектных файлах отсутствуют. "
        "В связи с этим в отчете используется формулировка о прохождении практики в рамках учебного проекта "
        "на базе собственного рабочего места студента и серверной среды, подготовленной для размещения "
        "веб-приложения. Подразделение практики условно определяется как учебная проектная среда веб-разработки, "
        "в которой выполнялись анализ требований, проектирование, программирование, тестирование, отладка и "
        "подготовка проекта к развертыванию.",
    )
    p(
        doc,
        "Направление деятельности в рамках практики соответствовало профилю веб-разработчика: создание клиентской "
        "части приложения, разработка серверного REST API, настройка обмена данными в реальном времени через "
        "WebSocket, проектирование базы данных, настройка авторизации, работа с безопасностью и подготовка "
        "деплойных конфигураций для VPS-сервера.",
    )
    p(
        doc,
        "В процессе разработки использовались технические и программные средства, фактически обнаруженные в проекте: "
        "Node.js, npm workspaces, TypeScript, React, Vite, Tailwind CSS, Express, Socket.IO, MySQL, Nginx, systemd, "
        "Git, Service Worker, а также средства автоматизированной проверки: TypeScript compiler, Vite production "
        "build и тесты Node.js через tsx --test.",
    )
    screenshot_marker(doc, "главный экран мессенджера «БренксЧат» со списком чатов и открытой перепиской")

    h(doc, "1 Цели и задачи практики", 1)
    p(
        doc,
        "Целью производственной практики являлось комплексное освоение практических навыков веб-разработки путем "
        "самостоятельного создания полнофункционального клиент-серверного приложения. Разработка мессенджера "
        "позволила применить знания по программированию, проектированию интерфейсов, организации сетевого "
        "взаимодействия, работе с базами данных, обеспечению безопасности и сопровождению приложения.",
    )
    p(doc, "В ходе практики были поставлены следующие задачи:")
    for item in [
        "изучить предметную область мессенджеров и определить состав пользовательских сценариев;",
        "спроектировать клиент-серверную архитектуру приложения;",
        "разработать структуру данных для пользователей, чатов, сообщений, сессий, push-подписок и E2EE-устройств;",
        "реализовать регистрацию, вход, подтверждение почты, сброс пароля и режим «запомнить меня»;",
        "реализовать обмен сообщениями в реальном времени через Socket.IO;",
        "реализовать личные чаты, группы, каналы, избранное, поиск, ответы, реакции, редактирование, удаление и пересылку сообщений;",
        "добавить отправку изображений, файлов, голосовых сообщений и видеокружков;",
        "реализовать голосовые и видеозвонки на базе WebRTC-сигналинга;",
        "настроить хранение данных в MySQL и подготовить конфигурации для серверного развертывания;",
        "провести тестирование, исправить ошибки и подготовить проектную документацию.",
    ]:
        bullet(doc, item)
    p(
        doc,
        "В результате выполнения работы были сформированы профессиональные компетенции, связанные с разработкой "
        "программного кода, проектированием интерфейсов, взаимодействием с серверной частью, использованием "
        "баз данных, анализом ошибок, тестированием и сопровождением веб-приложения. Выполненная работа напрямую "
        "связана с профессиональной деятельностью веб-разработчика, так как охватывает полный цикл создания "
        "современного веб-сервиса: от структуры проекта и дизайна интерфейса до серверного API, БД, безопасности "
        "и деплоя.",
    )

    h(doc, "2 База практики", 1)
    p(
        doc,
        "База практики представлена учебной проектной средой, включающей локальное рабочее место разработчика, "
        "репозиторий исходного кода и серверную инфраструктуру для размещения веб-приложения. Точные характеристики "
        "аппаратного обеспечения в файлах проекта не зафиксированы. По структуре путей рабочей директории "
        "(/Users/roz1es/...) можно сделать вывод, что разработка велась в macOS-совместимой локальной среде. "
        "Для развертывания подготовлены конфигурации Ubuntu/VPS: Nginx, systemd, MySQL и резервное копирование.",
    )
    p(
        doc,
        "Аппаратное обеспечение, которое необходимо указать при финальном заполнении отчета: персональный компьютер "
        "или ноутбук разработчика, монитор, клавиатура, мышь, сетевой адаптер и доступ к интернету. Для серверной "
        "части использовался VPS с Linux/Ubuntu, что подтверждается конфигурациями deploy/systemd/brenkschat.service "
        "и deploy/nginx/silentx.ru.conf. Точные параметры CPU, RAM и диска VPS в проектных файлах не указаны.",
    )
    tool_rows = [
        ["Язык программирования", "TypeScript", "Основной язык клиентской и серверной разработки; используется в client/src и server/src."],
        ["Среда выполнения", "Node.js", "Запуск серверной части, сборочных скриптов, тестов и production-приложения."],
        ["Менеджер пакетов", "npm workspaces", "Управление зависимостями client и server из корневого package.json."],
        ["Клиентский фреймворк", "React 18", "Построение интерфейса мессенджера, страниц входа, чатов и модальных окон."],
        ["Маршрутизация", "React Router", "Разделение маршрутов /login, / и /admin в client/src/App.tsx."],
        ["Сборщик", "Vite", "Локальный dev-сервер на порту 5173, proxy /api и /socket.io, production-сборка клиента."],
        ["CSS-фреймворк", "Tailwind CSS", "Стилизация интерфейса, темы, адаптивность, анимации сообщений и модальных окон."],
        ["Серверный фреймворк", "Express", "REST API для авторизации, чатов, пользователей, админ-панели, push и E2EE."],
        ["Realtime", "Socket.IO", "Обмен сообщениями, typing, presence, реакции, звонки и события чатов в реальном времени."],
        ["База данных", "MySQL", "Постоянное хранение пользователей, чатов, сообщений, сессий и служебных данных."],
        ["Драйвер БД", "mysql2", "Подключение Node.js-сервера к MySQL через DATABASE_URL."],
        ["Авторизация", "jsonwebtoken", "Формирование JWT с идентификатором пользователя и сессии."],
        ["Хэширование паролей", "argon2", "Безопасное хранение паролей с алгоритмом Argon2id."],
        ["Почта", "Nodemailer / Resend API", "Отправка кодов подтверждения входа, регистрации, привязки почты и сброса пароля."],
        ["Push-уведомления", "web-push, Service Worker", "Уведомления о новых сообщениях и входящих звонках."],
        ["Шифрование", "libsodium-wrappers, Web Crypto API", "E2EE-текстовые сообщения, Curve25519-ключи, резервная копия ключей AES-GCM/PBKDF2."],
        ["Звонки", "WebRTC", "Голосовые и видеозвонки; ICE/STUN/TURN задаются через API /api/calls/ice-servers."],
        ["Веб-сервер", "Nginx", "HTTPS reverse proxy для API, Socket.IO, статических файлов и phpMyAdmin."],
        ["Сервисная среда", "systemd", "Автозапуск и hardening Node.js-сервера на VPS."],
        ["Контроль версий", "Git", "История коммитов проекта; в журнале есть начальные коммиты и конфигурация Railway."],
        ["Тестирование", "tsx --test, TypeScript compiler, Vite build", "Автоматизированные серверные тесты, проверка типов, production-сборка."],
        ["Документация", "README.md, DESKTOP_RELEASE.md, deploy/*.conf", "Описание запуска, структуры, desktop-направления и серверного развертывания."],
        ["Дополнительный desktop-клиент", "Flutter / Dart", "В каталоге desktop находится нативный MVP-клиент для macOS/Windows, не являющийся webview-оберткой."],
        ["Пакетирование", "Electron / electron-builder", "В корневом package.json сохранены скрипты сборки desktop-упаковки для macOS/Windows/Linux."],
    ]
    table(doc, "Таблица 2 — Используемые технические и программные средства", ["Тип средства", "Название", "Назначение"], tool_rows, [4.2, 4.6, 8.7])
    p(
        doc,
        "Структура вычислительной среды включает локальный клиент, сервер Node.js, СУБД MySQL и внешний reverse proxy. "
        "В режиме разработки клиент запускается Vite на порту 5173, а API и Socket.IO проксируются на сервер. "
        "В production Nginx обслуживает статические файлы и перенаправляет API/WebSocket-запросы на Node.js-сервис. "
        "Для проектирования использовались TypeScript-интерфейсы в server/src/types.ts и client/src/types/index.ts, "
        "структура каталогов компонентов и контекстов, а также SQL-схема, создаваемая в persist.ts, sessions.ts и e2eeDevices.ts.",
    )

    h(doc, "3 Содержание практики", 1)
    p(
        doc,
        "Содержание практики включало последовательную разработку мессенджера, начиная с анализа требований и "
        "заканчивая тестированием, документированием и подготовкой инфраструктуры для развертывания. Ниже приведено "
        "описание основных этапов с указанием использованных технологий, подтверждающих файлов и результата.",
    )

    stage_rows = [
        [
            "Анализ требований",
            "Определены основные сценарии: регистрация, вход, подтверждение почты, личные чаты, группы, каналы, медиа, звонки, администрирование.",
            "README.md, client/src/pages/LoginPage.tsx, client/src/pages/MessengerPage.tsx, client/src/pages/AdminPage.tsx",
            "Сформирован состав функций мессенджера и пользовательских ролей.",
        ],
        [
            "Проектирование системы",
            "Выбрана архитектура SPA + REST API + WebSocket. Клиент разделен на страницы, компоненты и контексты; сервер — на API, realtime-обработчики, store и persistence.",
            "client/src/App.tsx, client/src/contexts/*.tsx, server/src/index.ts, server/src/socketHandlers.ts, server/src/store.ts",
            "Получена расширяемая структура проекта с отдельными зонами ответственности.",
        ],
        [
            "Проектирование БД",
            "Спроектированы таблицы users, chats, messages, chat_participants, auth_sessions, push_subscriptions, E2EE-таблицы и служебные таблицы.",
            "server/src/persist.ts, server/src/sessions.ts, server/src/e2eeDevices.ts, server/src/types.ts",
            "Создана MySQL-схема для постоянного хранения данных мессенджера.",
        ],
        [
            "Разработка серверной части",
            "Реализованы REST-эндпоинты Express, Socket.IO-события, проверка прав, рассылка уведомлений, нормализация медиа и обработка звонков.",
            "server/src/index.ts, server/src/socketHandlers.ts, server/src/pushNotifications.ts, server/src/email.ts",
            "Сервер обеспечивает API, realtime-обмен, авторизацию и интеграцию с инфраструктурой.",
        ],
        [
            "Разработка клиентской части",
            "Созданы страницы входа, мессенджера и админ-панели; реализованы компоненты чата, строки ввода, профиля, просмотра фото, тем и модальных окон.",
            "client/src/pages/*.tsx, client/src/components/*.tsx, client/src/index.css, client/tailwind.config.js",
            "Получен адаптивный пользовательский интерфейс с темной и светлой темой.",
        ],
        [
            "Авторизация и аутентификация",
            "Добавлены регистрация, вход, подтверждение кода из письма, сброс пароля, cookie-сессии, JWT, режим remember me и роль администратора.",
            "server/src/auth.ts, server/src/password.ts, server/src/emailAuth.ts, server/src/sessions.ts, client/src/contexts/AuthContext.tsx",
            "Пользователь получает защищенную сессию, а администратор — доступ к отдельной панели.",
        ],
        [
            "Обмен сообщениями",
            "Реализованы отправка, редактирование, удаление, реакции, ответы, пересылка, закрепление, прочтение, поиск, typing и presence.",
            "server/src/socketHandlers.ts, client/src/contexts/MessengerContext.tsx, client/src/components/MessageBubble.tsx, client/src/components/MessageInput.tsx",
            "Сообщения доставляются в реальном времени и синхронизируются между участниками чата.",
        ],
        [
            "Работа с медиа и звонками",
            "Добавлены изображения, файлы, голосовые сообщения, видеокружки, WebRTC-звонки, уведомления о звонках и настройка ICE-серверов.",
            "client/src/components/VoiceVideoMedia.tsx, client/src/contexts/CallContext.tsx, client/src/components/PhotoViewer.tsx, server/src/index.ts",
            "Мессенджер поддерживает мультимедийные сценарии общения.",
        ],
        [
            "Работа с базой данных",
            "Реализованы загрузка и сохранение состояния в MySQL, миграция из app_state/JSON, периодическое сохранение, резервное копирование.",
            "server/src/persist.ts, deploy/backup/brenkschat-backup.mjs, server/scripts/*.mjs",
            "Данные приложения могут храниться на сервере и переноситься между средами.",
        ],
        [
            "Тестирование и отладка",
            "Проверены password/session/E2EE/store-сценарии, выполнены npm test -w server и npm run build.",
            "server/src/*.test.ts, package.json, server/package.json, client/package.json",
            "8 серверных тестов пройдено, production-сборка клиента и сервера выполнена успешно.",
        ],
        [
            "Документирование и деплой",
            "Подготовлены README, desktop-документация, systemd service, Nginx-конфигурация, резервное копирование и env-пример.",
            "README.md, DESKTOP_RELEASE.md, server/.env.example, deploy/nginx/silentx.ru.conf, deploy/systemd/brenkschat.service",
            "Проект имеет инструкции запуска и основу для размещения на VPS.",
        ],
    ]
    table(doc, "Таблица 3 — Этапы выполнения производственной практики", ["Этап", "Что было сделано", "Файлы проекта", "Результат"], stage_rows, [3.1, 5.2, 5.1, 4.1])

    h(doc, "3.1 Анализ требований", 2)
    p(
        doc,
        "На первом этапе была определена предметная область проекта. Мессенджер должен был обеспечивать быстрое "
        "общение пользователей, поддерживать персональные и групповые сценарии, иметь современный интерфейс, "
        "работать в браузере и поддерживать дальнейшее серверное развертывание. README.md фиксирует исходную "
        "идею Telegram-style мессенджера на React + TypeScript и Node.js + Express + Socket.IO. В процессе "
        "развития проекта функциональность была расширена: появились подтверждение почты, MySQL, E2EE-слой, "
        "звонки, push-уведомления, PWA-манифест и административная панель.",
    )
    screenshot_marker(doc, "страница входа с вкладками входа, регистрации, сброса пароля и кодом подтверждения")

    h(doc, "3.2 Проектирование системы", 2)
    p(
        doc,
        "Архитектура приложения построена по клиент-серверной модели. Клиентская часть находится в каталоге client "
        "и реализует интерфейс пользователя. Серверная часть находится в каталоге server и отвечает за API, "
        "обработку realtime-событий и сохранение данных. В client/src/App.tsx описаны маршруты /login, / и /admin. "
        "Страница MessengerPage объединяет CallProvider, ChatWallpaperProvider и MessengerProvider, а также основные "
        "компоненты интерфейса: список чатов, окно переписки, шапку чата, строку ввода и модальные окна.",
    )
    p(
        doc,
        "Серверная часть разделена на index.ts, где объявлены HTTP-эндпоинты Express, socketHandlers.ts, где "
        "обрабатываются события Socket.IO, store.ts, где описаны операции над состоянием, и persist.ts, где "
        "реализовано постоянное хранение данных. Такой подход позволяет отделить транспортный слой от бизнес-логики "
        "и от слоя работы с базой данных.",
    )

    h(doc, "3.3 Проектирование базы данных", 2)
    p(
        doc,
        "В проекте предусмотрено два варианта хранения: резервный файловый JSON и основной MySQL-режим, который "
        "включается при PERSIST_BACKEND=mysql и наличии DATABASE_URL. MySQL-схема создается программно при запуске "
        "через ensureMysqlSchema, initializeSessionStore и initializeE2eeDevices. Это снижает риск ручной ошибки "
        "при развертывании приложения.",
    )
    db_rows = [
        ["users", "Пользователи", "id, username, password, email, email_verified, avatar_url, display_name, bio, phone, birth_date, is_admin, banned, privacy"],
        ["chats", "Чаты", "id, type, name, avatar_url, last_message, unread, last_read_at, pinned_message_id, channel_owner_id"],
        ["chat_participants", "Связь пользователей и чатов", "chat_id, user_id, position; внешний ключ на chats(id) с ON DELETE CASCADE"],
        ["messages", "Сообщения", "id, chat_id, sender_id, text, encrypted_text, image_url, media_*, created_at, deleted, edited_at, reply_to_message_id, reactions"],
        ["user_chat_muted", "Отключенные уведомления", "user_id, chat_id"],
        ["user_chat_pinned", "Закрепление чатов в списке", "user_id, chat_id"],
        ["push_subscriptions", "Push-подписки браузеров", "user_id, endpoint, data"],
        ["auth_sessions", "Активные сессии", "id, user_id, created_at, expires_at"],
        ["user_e2ee_devices", "Публичные ключи E2EE-устройств", "user_id, device_id, public_key, created_at, last_seen_at"],
        ["user_e2ee_key_backups", "Резервные копии ключей E2EE", "user_id, version, salt, iv, ciphertext, iterations, updated_at"],
        ["app_state", "Совместимость со старым форматом", "id, data, updated_at"],
    ]
    table(doc, "Таблица 4 — Структура базы данных проекта", ["Таблица", "Назначение", "Основные поля"], db_rows, [3.7, 4.4, 9.4])
    screenshot_marker(doc, "фрагмент phpMyAdmin или схемы MySQL с таблицами users, chats и messages")

    h(doc, "3.4 Разработка серверной части", 2)
    p(
        doc,
        "Серверная часть реализована на Express и Socket.IO. В server/src/index.ts объявлены маршруты для регистрации, "
        "входа, подтверждения входа кодом, выхода, сброса пароля, управления профилем, поиска пользователей, "
        "создания чатов, групп и каналов, получения сообщений, работы с E2EE-устройствами, push-подписками и "
        "админ-панелью. В server/src/socketHandlers.ts реализованы события send_message, edit_message, "
        "delete_message, toggle_reaction, forward_messages, mark_read, typing и call_signal.",
    )
    p(
        doc,
        "Для защиты API применяются Helmet, CORS, rate limiting и проверка авторизации. Конфигурация server/.env.example "
        "предусматривает CLIENT_ORIGINS, ключи VAPID, провайдера почты и параметры SMTP/Resend. В production-среде "
        "Nginx проксирует /api и /socket.io на Node.js-сервер, а также обслуживает статические файлы клиентской сборки.",
    )

    h(doc, "3.5 Разработка клиентской части", 2)
    p(
        doc,
        "Клиентская часть реализована как одностраничное React-приложение. Компонентная структура позволяет разделить "
        "ответственность между отдельными элементами интерфейса. ChatList отвечает за список чатов и поиск, ChatWindow "
        "отображает переписку, MessageBubble отвечает за внешний вид сообщения и действия с ним, MessageInput реализует "
        "ввод текста, вложения, вставку изображений из буфера, запись голосовых сообщений и видеокружков. PhotoViewer "
        "используется для просмотра изображений, VoiceVideoMedia — для голосовых и видео-сообщений.",
    )
    p(
        doc,
        "Для управления состоянием используются React Context: AuthContext хранит текущего пользователя и методы "
        "авторизации; MessengerContext синхронизирует чаты, сообщения, typing, реакции и Socket.IO; CallContext "
        "управляет голосовыми и видеозвонками; ThemeContext отвечает за светлую и темную тему; ChatWallpaperContext "
        "хранит настройки фонов чата. Стилизация реализована Tailwind CSS и дополнительными CSS-классами в client/src/index.css.",
    )
    screenshot_marker(doc, "модальное окно создания нового чата, группы или канала в стеклянном стиле")

    h(doc, "3.6 Авторизация и аутентификация", 2)
    p(
        doc,
        "Регистрация и вход реализованы с учетом подтверждения электронной почты. На сервере emailAuth.ts создает "
        "одноразовые коды, а email.ts отправляет письма через Resend, SMTP или выводит письмо в консоль в режиме "
        "разработки. Пароли хранятся не в открытом виде: password.ts использует Argon2id, проверяет минимальную "
        "длину и поддерживает миграцию устаревших учетных данных.",
    )
    p(
        doc,
        "После успешного входа сервер выдает JWT, связанный с записью в auth_sessions. Сессии могут быть короткими "
        "или длительными, если пользователь выбрал режим «запомнить меня». В auth.ts проверяется cookie brenks_session "
        "или Bearer-токен, наличие активной сессии, существование пользователя и отсутствие блокировки. Для администратора "
        "предусмотрена дополнительная проверка requireAdmin.",
    )

    h(doc, "3.7 Реализация обмена сообщениями", 2)
    p(
        doc,
        "Обмен сообщениями построен на Socket.IO. При подключении пользователь присоединяется к комнатам своих чатов "
        "и комнате user:<id>. При отправке сообщения сервер проверяет участие пользователя в чате, права отправки в "
        "канал, корректность зашифрованного текста и параметры медиа. После сохранения сообщение рассылается участникам, "
        "обновляется preview последнего сообщения, счетчики непрочитанных и push-уведомления.",
    )
    p(
        doc,
        "Для удобства общения реализованы ответы на сообщения, реакции, пересылка, редактирование, удаление, закрепление, "
        "мьют чата, закрепление чата в списке, отметка прочтения и индикатор печати. В личных чатах предусмотрен слой "
        "E2EE для текста: клиент формирует envelope с алгоритмом crypto_box_curve25519xsalsa20poly1305, а сервер хранит "
        "его как JSON без расшифровки текста.",
    )
    screenshot_marker(doc, "пример переписки с текстовым сообщением, реакциями, ответом и медиавложением")

    h(doc, "3.8 Работа с базой данных", 2)
    p(
        doc,
        "Слой persist.ts преобразует внутреннее состояние приложения в нормализованные таблицы MySQL и обратно. При "
        "чтении состояния сервер загружает пользователей, чаты, участников, сообщения, настройки mute/pin и push-подписки. "
        "При сохранении состояние записывается в соответствующие таблицы. Для совместимости сохранена таблица app_state, "
        "которая содержит JSON-снимок старого формата.",
    )
    p(
        doc,
        "Для эксплуатации подготовлен backup-скрипт deploy/backup/brenkschat-backup.mjs. Он выполняет mysqldump с "
        "параметрами безопасного дампа, архивирует исходный код без node_modules, client/dist, server/dist, server/data "
        "и .git, сохраняет важные конфигурационные файлы, формирует SHA256SUMS и metadata.json, а также удаляет старые "
        "резервные копии старше установленного срока хранения.",
    )

    h(doc, "3.9 Тестирование и исправление ошибок", 2)
    p(
        doc,
        "В проекте обнаружены автоматизированные тесты server/src/e2eeDevices.test.ts, server/src/password.test.ts, "
        "server/src/sessions.test.ts и server/src/store.security.test.ts. Они проверяют регистрацию и неизменяемость "
        "Curve25519-ключей устройств, допустимость резервных копий E2EE-ключей, хэширование и проверку паролей Argon2id, "
        "поведение устаревших паролей, запрет отключенных аккаунтов, изоляцию сессий, ограничение длительности remember-me "
        "сессии и миграцию учетных данных без сброса администраторских паролей.",
    )
    p(
        doc,
        "Команда npm test -w server была выполнена успешно: 8 тестов пройдено, 0 тестов завершились ошибкой. Команда "
        "npm run build также выполнена успешно: клиент собран Vite, сервер скомпилирован TypeScript-компилятором. "
        "Автоматизированных React-тестов в проекте не обнаружено; для клиентской части применялась ручная проверка "
        "через браузер, локальный dev-сервер и production-сборку.",
    )
    p(
        doc,
        "В ходе отладки устранялись ошибки, связанные с отображением интерфейса, работой тем, контекстным меню, удалением "
        "сообщений, показом медиа, записью голосовых сообщений и видеокружков, подтверждением почты, настройкой DNS/SSL "
        "и переносом состояния в MySQL. Наличие миграционных скриптов migrate-video-notes.mjs, migrate-voice-messages.mjs "
        "и remove-users-without-email.mjs подтверждает, что проект развивался итерационно и требовал сопровождения данных.",
    )

    h(doc, "3.10 Документирование проекта", 2)
    p(
        doc,
        "Документация проекта представлена README.md, где описана структура проекта, запуск в режиме разработки, "
        "production-сборка и основные функции. В DESKTOP_RELEASE.md и desktop/README.md зафиксированы материалы по "
        "desktop-направлению. В deploy находятся конфигурации для Nginx, systemd, SSH hardening и резервного копирования. "
        "Файл server/.env.example описывает переменные окружения для почты, CORS, push-уведомлений и других параметров.",
    )

    h(doc, "3.11 Список выполненных работ", 2)
    completed_items = [
        "Проведен анализ предметной области и требований к веб-мессенджеру.",
        "Создана структура монорепозитория с рабочими областями client и server.",
        "Разработан интерфейс входа, регистрации, сброса пароля и подтверждения кода из письма.",
        "Реализован интерфейс списка чатов, окна переписки, профилей, модальных окон и просмотра изображений.",
        "Настроен обмен сообщениями в реальном времени через Socket.IO.",
        "Реализованы личные чаты, группы, каналы и избранное.",
        "Добавлены реакции, ответы, пересылка, редактирование, удаление, закрепление сообщений и индикатор печати.",
        "Добавлена отправка изображений, файлов, голосовых сообщений и видеокружков.",
        "Реализованы голосовые и видеозвонки с использованием WebRTC-сигналинга.",
        "Реализованы push-уведомления через Service Worker и web-push.",
        "Настроено хранение данных в MySQL и служебные таблицы для сессий и E2EE.",
        "Реализованы меры безопасности: Argon2id, JWT-сессии, rate limiting, Helmet, email-подтверждение, админ-права.",
        "Подготовлены конфигурации Nginx, systemd и резервного копирования.",
        "Проведена сборка проекта и запуск автоматизированных серверных тестов.",
        "Подготовлен официальный отчет по производственной практике на основе исходного кода проекта.",
    ]
    for idx, item in enumerate(completed_items, 1):
        manual_numbered(doc, idx, item)

    h(doc, "4 Охрана труда", 1)
    p(
        doc,
        "При выполнении работ по разработке веб-приложения основным рабочим инструментом являлся персональный компьютер. "
        "Работа за компьютером относится к видам деятельности, требующим соблюдения требований эргономики, электробезопасности "
        "и рационального режима труда и отдыха. Несоблюдение данных требований может привести к повышенной утомляемости, "
        "перенапряжению зрения, нарушению осанки и снижению работоспособности.",
    )
    h(doc, "4.1 Требования к рабочему месту", 2)
    p(
        doc,
        "Рабочее место должно обеспечивать удобное положение пользователя. Стол должен иметь достаточную площадь для "
        "монитора, клавиатуры, мыши и документов. Кресло должно регулироваться по высоте и поддерживать спину. Монитор "
        "следует размещать на расстоянии примерно 50–70 см от глаз, верхняя граница экрана должна находиться на уровне "
        "глаз или немного ниже. Клавиатура и мышь должны располагаться так, чтобы кисти рук не находились в постоянном "
        "напряжении.",
    )
    h(doc, "4.2 Требования к освещению", 2)
    p(
        doc,
        "Освещение должно быть равномерным, без резких бликов на экране. Рабочее место рекомендуется располагать так, "
        "чтобы естественный свет падал сбоку. При недостаточном освещении необходимо использовать искусственный источник "
        "света, не направленный прямо в глаза пользователя и не создающий отражений на мониторе.",
    )
    h(doc, "4.3 Режим труда и отдыха", 2)
    p(
        doc,
        "При длительной работе за компьютером необходимо делать регулярные перерывы. Рекомендуется каждые 45–60 минут "
        "выполнять короткий перерыв продолжительностью 5–10 минут, а также периодически выполнять упражнения для глаз "
        "и кистей рук. Во время практики особенно важно чередовать программирование, чтение документации, тестирование "
        "и оформление отчета, чтобы снизить монотонную нагрузку.",
    )
    h(doc, "4.4 Электробезопасность", 2)
    p(
        doc,
        "Перед началом работы необходимо убедиться в исправности кабелей питания, розеток, сетевых фильтров и оборудования. "
        "Запрещается работать с поврежденными проводами, разбирать оборудование под напряжением, ставить жидкости рядом "
        "с системным блоком, ноутбуком или периферией. Подключение оборудования должно выполняться только сухими руками. "
        "При появлении запаха гари, искрения или нестабильной работы устройства следует прекратить работу и отключить питание.",
    )
    h(doc, "4.5 Правила работы с компьютерной техникой", 2)
    p(
        doc,
        "Не допускается установка неизвестного программного обеспечения без проверки источника. При работе с серверными "
        "учетными данными, ключами почтовых сервисов, JWT_SECRET, DATABASE_URL и другими секретами необходимо исключать "
        "их публикацию в открытом репозитории. Резервные копии следует хранить в защищенной директории с ограниченными "
        "правами доступа. В проекте эти требования частично отражены через server/.env.example, systemd UMask=0077 и "
        "backup-скрипт, создающий файлы с правами 0600/0700.",
    )
    h(doc, "4.6 Профилактика утомления", 2)
    p(
        doc,
        "Для профилактики утомления рекомендуется соблюдать правильную посадку, регулярно проветривать помещение, "
        "использовать комфортную яркость экрана, избегать длительной работы в полной темноте, выполнять гимнастику "
        "для глаз и разминку мышц шеи, плеч и кистей. При разработке интерфейса также следует учитывать зрительную "
        "нагрузку конечных пользователей: достаточный контраст текста, читаемый размер шрифта, понятные состояния "
        "кнопок и отсутствие резких мигающих эффектов.",
    )

    h(doc, "Вывод", 1)
    p(
        doc,
        "В ходе производственной практики была выполнена разработка полнофункционального веб-мессенджера «БренксЧат». "
        "Проект включает клиентскую часть на React и TypeScript, серверную часть на Node.js, Express и Socket.IO, "
        "постоянное хранение данных в MySQL, авторизацию с подтверждением почты, работу с сообщениями в реальном "
        "времени, медиафайлами, голосовыми сообщениями, видеокружками, WebRTC-звонками, push-уведомлениями и "
        "административными функциями.",
    )
    p(
        doc,
        "В процессе практики были закреплены навыки проектирования структуры веб-приложения, разработки REST API, "
        "использования WebSocket, работы с MySQL, настройки сессий и ролей, реализации безопасного хранения паролей, "
        "интеграции почтовых уведомлений, настройки production-сборки и серверного развертывания. Особое значение "
        "имело исправление ошибок и итерационная доработка проекта, так как мессенджер содержит множество взаимосвязанных "
        "частей: интерфейс, realtime-обмен, хранение данных, авторизацию, медиа и звонки.",
    )
    p(
        doc,
        "Разработка мессенджера позволила получить практический опыт, максимально приближенный к реальной деятельности "
        "веб-разработчика. Были сформированы компетенции по анализу требований, проектированию архитектуры, написанию "
        "клиентского и серверного кода, работе с базами данных, тестированию, отладке, документированию и обеспечению "
        "безопасности веб-приложений. Поставленные задачи практики выполнены, а полученный результат может быть использован "
        "как основа для дальнейшего развития проекта: масштабирования БД, улучшения UI/UX, расширения E2EE на медиа, "
        "повышения надежности звонков через TURN и выпуска desktop-клиента.",
    )

    h(doc, "Список использованных источников", 1)
    sources = [
        "Исходные материалы проекта «БренксЧат»: README.md, package.json, client/src, server/src, deploy, desktop. — Локальный репозиторий проекта, 2026.",
        "React Documentation [Электронный ресурс]. — Режим доступа: https://react.dev/ (дата обращения: 16.06.2026).",
        "TypeScript Documentation [Электронный ресурс]. — Режим доступа: https://www.typescriptlang.org/docs/ (дата обращения: 16.06.2026).",
        "Vite Guide [Электронный ресурс]. — Режим доступа: https://vite.dev/guide/ (дата обращения: 16.06.2026).",
        "Tailwind CSS Documentation [Электронный ресурс]. — Режим доступа: https://tailwindcss.com/docs (дата обращения: 16.06.2026).",
        "Express Documentation [Электронный ресурс]. — Режим доступа: https://expressjs.com/ (дата обращения: 16.06.2026).",
        "Socket.IO Documentation [Электронный ресурс]. — Режим доступа: https://socket.io/docs/ (дата обращения: 16.06.2026).",
        "Node.js Documentation [Электронный ресурс]. — Режим доступа: https://nodejs.org/docs/ (дата обращения: 16.06.2026).",
        "MySQL Reference Manual [Электронный ресурс]. — Режим доступа: https://dev.mysql.com/doc/ (дата обращения: 16.06.2026).",
        "Nginx Documentation [Электронный ресурс]. — Режим доступа: https://nginx.org/en/docs/ (дата обращения: 16.06.2026).",
        "MDN Web Docs: WebRTC API [Электронный ресурс]. — Режим доступа: https://developer.mozilla.org/docs/Web/API/WebRTC_API (дата обращения: 16.06.2026).",
        "MDN Web Docs: Service Worker API [Электронный ресурс]. — Режим доступа: https://developer.mozilla.org/docs/Web/API/Service_Worker_API (дата обращения: 16.06.2026).",
        "OWASP Cheat Sheet Series: Password Storage Cheat Sheet [Электронный ресурс]. — Режим доступа: https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html (дата обращения: 16.06.2026).",
        "libsodium Documentation [Электронный ресурс]. — Режим доступа: https://doc.libsodium.org/ (дата обращения: 16.06.2026).",
        "RFC 7519. JSON Web Token (JWT) [Электронный ресурс]. — Режим доступа: https://www.rfc-editor.org/rfc/rfc7519 (дата обращения: 16.06.2026).",
    ]
    for idx, source in enumerate(sources, 1):
        manual_numbered(doc, idx, source, left=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
