from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path("/Users/roz1es/Downloads/Дневник практики ПП 04.01(ВТ).docx")
OUT = Path("/Users/roz1es/Downloads/Дневник практики ПП 04.01 БренксЧат готовый.docx")
FONT = "Times New Roman"


def set_run_font(run, size: float = 12, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_text(
    paragraph,
    text: str,
    size: float = 12,
    bold: bool = False,
    align=None,
    first_line: float | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)
    if align is not None:
        paragraph.alignment = align
    if first_line is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line)
    paragraph.paragraph_format.line_spacing = 1.15


def set_cell_text(
    cell,
    text: str,
    size: float = 11,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(text.split("\n")):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size, bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def remove_row(table, row_index: int) -> None:
    table._tbl.remove(table.rows[row_index]._tr)


def clone_row(row):
    return deepcopy(row._tr)


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = tbl_cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fill_main_fields(doc: Document) -> None:
    set_paragraph_text(
        doc.paragraphs[5],
        "ПМ.04 Выполнение работ по разработке, тестированию и сопровождению веб-приложения",
        12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_paragraph_text(doc.paragraphs[7], "Студент ____ курса", 12)
    set_paragraph_text(doc.paragraphs[11], "Группа ВТ-Х-ХХ", 12)

    set_paragraph_text(
        doc.paragraphs[32],
        "Учебный проект по разработке веб-мессенджера «БренксЧат»",
        12,
    )
    set_paragraph_text(
        doc.paragraphs[34],
        "самостоятельная проектная работа в рамках производственной практики",
        12,
    )
    set_paragraph_text(
        doc.paragraphs[36],
        "Адрес профильной организации в исходных материалах не указан;",
        12,
    )
    set_paragraph_text(
        doc.paragraphs[37],
        "заполняется студентом перед сдачей при наличии официальной базы практики.",
        12,
    )

    set_paragraph_text(
        doc.paragraphs[40],
        "Практика имеет целью закрепление профессиональных навыков по специальности "
        "09.02.09 «Веб-разработка», формирование общих и профессиональных компетенций, "
        "а также приобретение практического опыта разработки, тестирования, отладки "
        "и документирования веб-приложения.",
        12,
        first_line=1.25,
    )
    set_paragraph_text(
        doc.paragraphs[72],
        "В результате изучения профессионального модуля ПМ.04 студент должен обладать "
        "профессиональными компетенциями, связанными с анализом требований, разработкой "
        "клиентской и серверной частей веб-приложения, работой с базой данных, тестированием, "
        "отладкой и подготовкой технической документации.",
        12,
        first_line=1.25,
    )
    competence_lines = [
        "- ПК 4.1. Выполнять проектирование и разработку клиентской части веб-приложения;",
        "- ПК 4.2. Выполнять проектирование и разработку серверной части веб-приложения и API;",
        "- ПК 4.3. Выполнять тестирование, отладку, документирование и подготовку веб-приложения к развертыванию.",
    ]
    for idx, line in enumerate(competence_lines, start=73):
        set_paragraph_text(doc.paragraphs[idx], line, 12, first_line=0)

    set_paragraph_text(
        doc.paragraphs[110],
        "Разработать веб-мессенджер «БренксЧат»: выполнить анализ требований, "
        "спроектировать клиент-серверную архитектуру и структуру базы данных, "
        "реализовать авторизацию, обмен сообщениями в реальном времени, работу с медиа, "
        "голосовые и видеосообщения, WebRTC-звонки, администрирование пользователей, "
        "сохранение данных в MySQL, тестирование, отладку, подготовку проекта к деплою "
        "и оформление отчетной документации.",
        12,
        first_line=1.25,
    )

    set_paragraph_text(
        doc.paragraphs[132],
        "анализ требований к мессенджеру; проектирование архитектуры SPA + REST API + WebSocket; "
        "разработка клиентской части на React и TypeScript; разработка серверной части на Node.js, "
        "Express и Socket.IO; проектирование MySQL-хранилища; реализация авторизации, чатов, "
        "сообщений, медиа, звонков, push-уведомлений, E2EE-механизмов; тестирование, исправление "
        "ошибок, деплой на VPS и подготовка документации.",
        12,
        first_line=1.25,
    )
    set_paragraph_text(
        doc.paragraphs[137],
        "зарекомендовал себя как ответственный исполнитель, способный самостоятельно анализировать "
        "требования, подбирать инструменты разработки, реализовывать клиентскую и серверную логику, "
        "работать с базой данных, выполнять отладку и оформлять результаты работы. В ходе практики "
        "продемонстрированы навыки веб-разработки, внимательность к качеству интерфейса и готовность "
        "исправлять выявленные ошибки.",
        12,
        first_line=1.25,
    )
    set_paragraph_text(
        doc.paragraphs[147],
        "По итогам производственной практики студент выполнил задание в полном объеме: разработал "
        "веб-мессенджер «БренксЧат», реализовал основные пользовательские функции, подготовил серверную "
        "и клиентскую части, выполнил тестирование и оформил отчетные материалы. Результаты практики "
        "соответствуют направлению подготовки по специальности 09.02.09 «Веб-разработка».",
        12,
        first_line=1.25,
    )


def fill_diary_table(doc: Document) -> None:
    entries = [
        (
            "08.06.2026",
            "Ознакомление с заданием производственной практики. Анализ предметной области и требований к веб-мессенджеру «БренксЧат». Определение основных пользовательских сценариев: регистрация, вход, список чатов, переписка, медиа и администрирование.",
        ),
        (
            "09.06.2026",
            "Проектирование общей архитектуры приложения. Выделение клиентской части, серверной части, REST API, WebSocket-слоя и хранилища данных. Подготовка структуры проекта с рабочими областями client и server.",
        ),
        (
            "10.06.2026",
            "Проектирование структуры базы данных и модели состояния приложения. Определение сущностей users, chats, messages, chat_participants, auth_sessions, push_subscriptions, user_e2ee_devices и user_e2ee_key_backups.",
        ),
        (
            "11.06.2026",
            "Разработка серверной части на Node.js, Express и Socket.IO. Реализация базовых маршрутов API, подключение middleware, обработка realtime-событий, подготовка логики получения чатов и сообщений.",
        ),
        (
            "12.06.2026",
            "Разработка клиентской части на React, TypeScript, Vite и Tailwind CSS. Реализация страниц входа, основного окна мессенджера, списка чатов, окна переписки, модальных окон и переключения темы.",
        ),
        (
            "13.06.2026",
            "Реализация обмена сообщениями в реальном времени: отправка, получение, редактирование, удаление, реакции, ответы, пересылка, закрепление сообщений, статусы прочтения, typing-состояния и presence.",
        ),
        (
            "15.06.2026",
            "Разработка механизмов авторизации и безопасности: регистрация, вход, подтверждение почты кодом, сброс пароля, cookie-сессии, JWT, режим «запомнить меня», Argon2id-хэширование паролей и защита API.",
        ),
        (
            "16.06.2026",
            "Реализация работы с медиа: отправка изображений и файлов, голосовые сообщения, видеокружки, просмотр фотографий, профиль пользователя, редактирование аватара и адаптация интерфейса под разные темы.",
        ),
        (
            "17.06.2026",
            "Настройка голосовых и видеозвонков на основе WebRTC и Socket.IO-сигналинга. Реализация контекста звонков, интерфейса входящего вызова, кнопок микрофона, камеры и звука, проверка ICE-настроек.",
        ),
        (
            "18.06.2026",
            "Настройка постоянного хранения данных в MySQL. Подготовка таблиц, миграционных сценариев, проверка подключения через mysql2, работа с phpMyAdmin, настройка серверной конфигурации и переменных окружения.",
        ),
        (
            "19.06.2026",
            "Тестирование и отладка проекта. Запуск npm run build и npm test -w server, проверка TypeScript-компиляции, исправление ошибок интерфейса, работы сообщений, медиа, E2EE-ключей и мобильной версии.",
        ),
        (
            "20.06.2026",
            "Подготовка проекта к развертыванию на VPS: сборка client/dist и server/dist, настройка Nginx, systemd-сервиса, HTTPS, резервного копирования и проверка работы приложения на домене silentx.ru.",
        ),
        (
            "21.06.2026",
            "Итоговая проверка работоспособности веб-мессенджера. Подготовка отчета по производственной практике, аттестационного листа и дневника практики. Формирование выводов о полученных навыках и результатах.",
        ),
    ]

    table = doc.tables[1]
    set_table_cell_margins(table)
    repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        shade_cell(cell, "EDEDED")
    set_cell_text(table.cell(0, 0), "Дата или период выполнения работ", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "Краткое содержание выполняемых работ", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 2), "Подпись руководителя практической подготовки от профильной организации", 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    # Оставляем только нужное количество строк, чтобы дневник не превращался в пачку пустых страниц.
    while len(table.rows) > len(entries) + 1:
        remove_row(table, len(table.rows) - 1)
    while len(table.rows) < len(entries) + 1:
        table._tbl.append(clone_row(table.rows[-1]))

    for row in table.rows:
        prevent_row_split(row)

    for idx, (date, work) in enumerate(entries, start=1):
        row = table.rows[idx]
        set_cell_text(row.cells[0], date, 10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], work, 9.5)
        set_cell_text(row.cells[2], "", 10, align=WD_ALIGN_PARAGRAPH.CENTER)


def configure_styles(doc: Document) -> None:
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            style.font.name = FONT
            if style.element.rPr is not None:
                style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def main() -> None:
    doc = Document(SRC)
    configure_styles(doc)
    fill_main_fields(doc)
    fill_diary_table(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
