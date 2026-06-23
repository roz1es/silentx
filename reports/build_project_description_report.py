from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/roz1es/Downloads/messengercursor 2/messengercursor")
OUT = Path("/Users/roz1es/Downloads/Техническое описание проекта БренксЧат.docx")

FONT = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "F2F4F7"
LIGHT_FILL = "F8FAFC"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size: float, bold: bool = False, color: str | None = None) -> None:
    style.font.name = FONT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        set_style_font(style, size, True, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        set_style_font(style, 11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "Техническое описание проекта «БренксЧат»"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], 9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    set_run_font(run, 9, color="666666")


def p(doc: Document, text: str = "", style: str | None = None, align=None):
    para = doc.add_paragraph(style=style)
    if text:
        run = para.add_run(text)
        set_run_font(run, 11)
    if align is not None:
        para.alignment = align
    return para


def heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    run = para.add_run(text)
    set_run_font(run, 11)
    return para


def number(doc: Document, text: str):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    run = para.add_run(text)
    set_run_font(run, 11)
    return para


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_grid(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_dxa):
                set_cell_width(cell, widths_dxa[idx])


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def write_cell(cell, text: str, bold: bool = False, size: float = 9.3, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(text.split("\n")):
        if idx:
            para.add_run().add_break()
        run = para.add_run(line)
        set_run_font(run, size, bold)


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    cap = p(doc, caption)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    for run in cap.runs:
        set_run_font(run, 10, True, DARK_BLUE)

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    set_table_grid(tbl, widths_dxa)
    set_cell_margins(tbl)
    repeat_table_header(tbl.rows[0])
    for idx, header in enumerate(headers):
        shade_cell(tbl.rows[0].cells[idx], HEADER_FILL)
        write_cell(tbl.rows[0].cells[idx], header, True, 9.2, WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            write_cell(cells[idx], value, False, 8.7)
        for cell in cells:
            if row and row[0].startswith("Итого"):
                shade_cell(cell, LIGHT_FILL)
        prevent_row_split(tbl.rows[-1])
    return tbl


def callout(doc: Document, title: str, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    set_table_grid(tbl, [TABLE_WIDTH_DXA])
    set_cell_margins(tbl, top=120, bottom=120, start=160, end=160)
    cell = tbl.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f"{title}. ")
    set_run_font(run, 10.5, True, DARK_BLUE)
    run = para.add_run(text)
    set_run_font(run, 10.5)


def add_title_page(doc: Document) -> None:
    p(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    title = p(doc, "Техническое описание проекта «БренксЧат»", align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.runs[0], 24, True, DARK_BLUE)

    subtitle = p(
        doc,
        "Веб-мессенджер: архитектура, технологии, база данных, хранимые данные и инфраструктура",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_run_font(subtitle.runs[0], 13, color="555555")

    meta_rows = [
        ["Проект", "БренксЧат"],
        ["Тип системы", "Веб-мессенджер с realtime-обменом сообщениями"],
        ["Основной стек", "React, TypeScript, Vite, Node.js, Express, Socket.IO, MySQL"],
        ["Дата подготовки", "20 июня 2026 г."],
    ]
    doc.add_paragraph()
    table(doc, "Краткая карточка проекта", ["Параметр", "Значение"], meta_rows, [2200, 7160])

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_overview(doc: Document) -> None:
    heading(doc, "1. Назначение проекта", 1)
    p(
        doc,
        "«БренксЧат» представляет собой веб-мессенджер для обмена сообщениями в реальном времени. "
        "Система поддерживает личные чаты, группы, каналы, медиа-сообщения, голосовые сообщения, "
        "видеокружки, реакции, ответы, пересылку, закрепление сообщений, статусы прочтения, поиск "
        "пользователей, профили, push-уведомления, голосовые и видеозвонки. Также реализована "
        "административная часть для просмотра состояния системы и управления пользователями.",
    )
    callout(
        doc,
        "Итог",
        "Проект является полноценным клиент-серверным приложением: клиент работает как SPA/PWA, "
        "сервер предоставляет REST API и Socket.IO-события, а постоянное состояние хранится в MySQL.",
    )

    heading(doc, "2. Общая архитектура", 1)
    p(
        doc,
        "Архитектура построена по модели SPA + REST API + WebSocket + MySQL. Клиентская часть "
        "отвечает за интерфейс и локальные браузерные возможности, серверная часть — за авторизацию, "
        "валидацию прав, работу с сообщениями, realtime-события, почтовые коды, push-уведомления и "
        "подключение к базе данных.",
    )
    for item in [
        "Клиент: одностраничное приложение React/TypeScript, собираемое Vite и стилизованное Tailwind CSS.",
        "Сервер: Node.js-приложение на Express и Socket.IO, запускаемое в production через systemd.",
        "База данных: MySQL-совместимое хранилище через библиотеку mysql2; в коде также сохранена совместимость с legacy JSON/app_state.",
        "Realtime: Socket.IO используется для доставки сообщений, реакций, typing-состояний, presence, обновлений чатов и сигналинга звонков.",
        "Инфраструктура: Nginx обслуживает статические файлы, проксирует /api и /socket.io на Node.js, HTTPS настроен через Certbot.",
    ]:
        bullet(doc, item)


def add_stack(doc: Document) -> None:
    heading(doc, "3. Используемые технологии и библиотеки", 1)
    rows = [
        ["Язык", "TypeScript", "Основной язык клиентской и серверной разработки."],
        ["Клиент", "React 18", "Компонентный пользовательский интерфейс мессенджера."],
        ["Маршрутизация", "React Router", "Маршруты входа, основного окна и админ-панели."],
        ["Сборка клиента", "Vite", "Dev-сервер, proxy /api и /socket.io, production-сборка client/dist."],
        ["Стили", "Tailwind CSS, PostCSS, Autoprefixer", "Темы, адаптивность, модальные окна, сообщения и анимации."],
        ["Realtime-клиент", "socket.io-client", "Подключение к Socket.IO-серверу из браузера."],
        ["Криптография клиента", "libsodium-wrappers, WebCrypto API", "E2EE для текстов личных сообщений и резервные копии ключей."],
        ["Сервер", "Node.js, Express", "REST API, middleware, отдача API и бизнес-логика."],
        ["Realtime-сервер", "Socket.IO", "Сообщения, реакции, presence, typing и WebRTC-сигналинг."],
        ["База данных", "MySQL + mysql2", "Постоянное хранение пользователей, чатов, сообщений и служебных данных."],
        ["Авторизация", "jsonwebtoken, cookie-сессии", "Сессии пользователя и проверка доступа к API."],
        ["Пароли", "argon2", "Хэширование паролей Argon2id."],
        ["Почта", "nodemailer", "Отправка кодов подтверждения, входа и восстановления пароля."],
        ["Безопасность API", "helmet, cors, express-rate-limit", "HTTP-заголовки безопасности, CORS и rate limit."],
        ["Push", "web-push, Service Worker", "Уведомления о сообщениях и звонках."],
        ["Desktop-артефакты", "Electron/electron-builder; отдельная папка Flutter desktop", "Подготовка к desktop-версии и сборщикам Windows/macOS."],
    ]
    table(doc, "Таблица 1 — Технологический стек проекта", ["Слой", "Средство", "Назначение"], rows, [1500, 2600, 5260])


def add_structure(doc: Document) -> None:
    heading(doc, "4. Структура проекта", 1)
    rows = [
        ["client/src/pages", "Страницы LoginPage, MessengerPage, AdminPage."],
        ["client/src/components", "Компоненты интерфейса: список чатов, окно чата, сообщение, строка ввода, профиль, просмотр фото, модальные окна."],
        ["client/src/contexts", "AuthContext, MessengerContext, CallContext, ThemeContext, ChatWallpaperContext."],
        ["client/src/lib", "API-клиент, E2EE, поиск пользователей, хранение auth-данных, обои чата, утилиты отображения."],
        ["client/public", "PWA manifest, service worker, логотипы и иконки."],
        ["server/src/index.ts", "Express-приложение, REST API, middleware, инициализация Socket.IO."],
        ["server/src/socketHandlers.ts", "Realtime-события сообщений, реакций, typing, прочтения, удаления, пересылки и звонков."],
        ["server/src/store.ts", "Бизнес-логика состояния: пользователи, чаты, сообщения, права, админ-обзор."],
        ["server/src/persist.ts", "MySQL-схема, чтение и сохранение нормализованного состояния."],
        ["server/src/sessions.ts", "Активные сессии пользователей."],
        ["server/src/e2eeDevices.ts", "E2EE-устройства и резервные копии ключей."],
        ["deploy", "Nginx, systemd, backup timer/service и SSH-hardening."],
        ["electron", "Electron-обертка для desktop-сборки."],
        ["desktop", "Черновая/отдельная Flutter desktop-реализация."],
    ]
    table(doc, "Таблица 2 — Основные каталоги и файлы", ["Путь", "Назначение"], rows, [2700, 6660])


def add_database(doc: Document) -> None:
    heading(doc, "5. База данных", 1)
    p(
        doc,
        "В проекте используется MySQL-совместимая база данных. Подключение выполняется через библиотеку "
        "mysql2 по строке DATABASE_URL. Основная схема нормализована и создается автоматически при запуске "
        "сервера. В коде также присутствует таблица app_state для совместимости со старым форматом хранения "
        "одного JSON-состояния, однако рабочая схема разделяет данные на отдельные таблицы.",
    )
    callout(
        doc,
        "Количество таблиц",
        "По фактическому коду проекта создается 11 таблиц: app_state, users, chats, chat_participants, "
        "messages, user_chat_muted, user_chat_pinned, push_subscriptions, auth_sessions, "
        "user_e2ee_devices, user_e2ee_key_backups.",
    )

    rows = [
        [
            "app_state",
            "Legacy-совместимость",
            "id, data, updated_at",
            "Единый JSON-снимок состояния старого формата. Используется как fallback при чтении старых данных.",
        ],
        [
            "users",
            "Пользователи",
            "id, username, password, email, email_verified, avatar_url, display_name, bio, phone, birth_date, is_admin, banned, privacy, updated_at",
            "Профили пользователей, хэши паролей, почта, аватар, отображаемое имя, описание, телефон, дата рождения, роль администратора, блокировка, privacy-настройки.",
        ],
        [
            "chats",
            "Чаты",
            "id, type, name, avatar_url, last_message, unread, last_read_at, pinned_message_id, channel_owner_id, updated_at",
            "Личные чаты, группы и каналы; аватар, последний preview, счетчики непрочитанных, прочтение, закрепленное сообщение, владелец канала.",
        ],
        [
            "chat_participants",
            "Участники чатов",
            "chat_id, user_id, position",
            "Связь many-to-many между пользователями и чатами. Есть внешний ключ на chats(id) с ON DELETE CASCADE.",
        ],
        [
            "messages",
            "Сообщения",
            "id, chat_id, sender_id, text, encrypted_text, image_url, media_kind, media_data_url, media_file_name, media_mime_type, media_duration_ms, created_at, deleted, edited_at, reply_to_message_id, reactions",
            "Текст/зашифрованный текст, медиа, файлы, голосовые, видеокружки, время, флаг удаления, редактирование, ответы и реакции.",
        ],
        [
            "user_chat_muted",
            "Отключение уведомлений",
            "user_id, chat_id",
            "Персональная настройка пользователя: чат без звука.",
        ],
        [
            "user_chat_pinned",
            "Закрепление чатов",
            "user_id, chat_id",
            "Персональная настройка пользователя: чат закреплен наверху списка.",
        ],
        [
            "push_subscriptions",
            "Push-уведомления",
            "user_id, endpoint, data",
            "Web Push-подписки браузеров: endpoint и JSON-ключи p256dh/auth.",
        ],
        [
            "auth_sessions",
            "Сессии входа",
            "id, user_id, created_at, expires_at",
            "Активные авторизованные сессии, включая длительные сессии remember me.",
        ],
        [
            "user_e2ee_devices",
            "E2EE-устройства",
            "user_id, device_id, public_key, created_at, last_seen_at",
            "Публичные Curve25519-ключи устройств пользователя для шифрования личных сообщений.",
        ],
        [
            "user_e2ee_key_backups",
            "Резервные копии E2EE-ключей",
            "user_id, version, salt, iv, ciphertext, iterations, updated_at",
            "Зашифрованная резервная копия приватного ключа устройства; шифруется паролем пользователя через WebCrypto/PBKDF2/AES-GCM.",
        ],
    ]
    table(doc, "Таблица 3 — Таблицы базы данных и хранимые данные", ["Таблица", "Назначение", "Основные поля", "Какие данные хранятся"], rows, [1650, 1700, 2950, 3060])

    heading(doc, "5.1. Связи и индексы", 2)
    for item in [
        "chat_participants.chat_id связан с chats.id внешним ключом; при удалении чата участники удаляются каскадно.",
        "messages.chat_id связан с chats.id внешним ключом; при удалении чата сообщения удаляются каскадно.",
        "messages имеет индекс idx_messages_chat_created по chat_id и created_at для быстрой загрузки истории чата.",
        "chat_participants имеет индекс idx_chat_participants_user для поиска чатов пользователя.",
        "auth_sessions имеет индексы по user_id и expires_at для очистки просроченных сессий.",
        "user_e2ee_devices имеет составной первичный ключ user_id + device_id.",
        "push_subscriptions имеет составной первичный ключ user_id + endpoint(191).",
    ]:
        bullet(doc, item)

    heading(doc, "5.2. Особенности хранения медиа", 2)
    p(
        doc,
        "Медиафайлы хранятся в таблице messages внутри полей media_* и image_url. По структуре проекта "
        "данные медиа передаются как data URL и сохраняются в LONGTEXT. Такой вариант удобен для учебного "
        "проекта и простого деплоя, но при росте объема данных логично вынести файлы в объектное хранилище "
        "или отдельное файловое хранилище, оставив в MySQL только ссылки и метаданные.",
    )


def add_features(doc: Document) -> None:
    heading(doc, "6. Реализованный функционал", 1)
    rows = [
        ["Аккаунты", "Регистрация, вход, подтверждение почты, сброс пароля, remember me, выход, профиль, аватар, privacy-настройки."],
        ["Чаты", "Личные чаты, избранное, группы, каналы, добавление участников, закрепление чатов, mute-режим."],
        ["Сообщения", "Отправка, редактирование, удаление, ответы, реакции, пересылка, закрепление, поиск, прочтение, typing."],
        ["Медиа", "Фотографии, файлы, голосовые сообщения, видеокружки, просмотр изображений, вставка/загрузка медиа."],
        ["Звонки", "WebRTC-аудио/видео звонки, сигналинг через Socket.IO, ICE-серверы, push-уведомления о входящем звонке."],
        ["Безопасность", "Argon2id, cookie/JWT-сессии, rate limit, Helmet, CORS, email-коды, E2EE для текстов личных сообщений."],
        ["Администрирование", "Админ-панель, обзор системы, блокировка пользователей, ссылка на базу данных/phpMyAdmin."],
        ["PWA", "manifest.json, service worker, push-уведомления, standalone-режим на телефоне."],
    ]
    table(doc, "Таблица 4 — Функциональные блоки", ["Блок", "Что реализовано"], rows, [1800, 7560])


def add_api_and_realtime(doc: Document) -> None:
    heading(doc, "7. API и realtime-взаимодействие", 1)
    p(
        doc,
        "Серверная часть предоставляет REST API для авторизации, пользователей, чатов, сообщений, "
        "E2EE-устройств, звонков, админ-панели и профиля. Для событий, которые должны приходить мгновенно, "
        "используется Socket.IO.",
    )
    api_rows = [
        ["Авторизация", "/api/register, /api/login, /api/login/confirm, /api/logout, /api/password-reset/*"],
        ["Профиль", "/api/me, /api/me/email/request, /api/me/email/confirm"],
        ["Пользователи", "/api/users/directory, /api/users/contacts, /api/users/search, /api/users/:userId"],
        ["Чаты", "/api/chats, /api/chats/:chatId/messages, /api/chats/direct, /api/chats/group, /api/chats/channel, /api/chats/saved"],
        ["Настройки чатов", "/api/chats/:chatId/mute, /pin-top, /pin-message, /clear, /members, PATCH /api/chats/:chatId, DELETE /api/chats/:chatId"],
        ["E2EE", "/api/e2ee/devices, /api/e2ee/key-backup, /api/chats/:chatId/e2ee-devices"],
        ["Звонки", "/api/calls/ice-servers"],
        ["Админ-панель", "/api/admin/overview, /api/admin/database, /api/admin/users/:userId/block"],
    ]
    table(doc, "Таблица 5 — Основные REST API-группы", ["Группа", "Маршруты"], api_rows, [2100, 7260])

    socket_rows = [
        ["join_chat", "Вход клиента в комнату чата."],
        ["send_message", "Отправка сообщения и рассылка участникам."],
        ["edit_message", "Редактирование сообщения."],
        ["delete_message", "Удаление сообщения."],
        ["toggle_reaction", "Добавление или снятие реакции."],
        ["forward_messages", "Пересылка сообщений в другой чат."],
        ["mark_read", "Отметка сообщений как прочитанных."],
        ["typing", "Индикатор набора текста."],
        ["presence", "Список пользователей онлайн с учетом privacy-настроек."],
        ["call_signal", "WebRTC-сигналинг: offer, answer, candidate, end и другие сигналы звонка."],
        ["chat_updated/message_edited/message_deleted/messages_cleared", "События синхронизации состояния интерфейса."],
    ]
    table(doc, "Таблица 6 — Основные Socket.IO-события", ["Событие", "Назначение"], socket_rows, [2600, 6760])


def add_security_and_deploy(doc: Document) -> None:
    heading(doc, "8. Безопасность", 1)
    for item in [
        "Пароли пользователей хранятся не в открытом виде, а как Argon2id-хэши.",
        "Авторизация работает через серверные сессии, связанные с cookie/JWT.",
        "Для входа, регистрации и сброса пароля используются email-коды.",
        "На авторизационные маршруты и API установлен express-rate-limit.",
        "Helmet добавляет защитные HTTP-заголовки; Nginx дополнительно задает HSTS, X-Frame-Options, X-Content-Type-Options, Referrer-Policy и CSP.",
        "Тексты личных сообщений могут шифроваться на клиенте: envelope хранится в messages.encrypted_text, а ключи устройств — в user_e2ee_devices.",
        "Резервная копия E2EE-ключа хранится зашифрованной и не содержит открытый приватный ключ.",
        "Production-сервис systemd работает от отдельного пользователя brenkschat и с hardening-настройками.",
    ]:
        bullet(doc, item)

    heading(doc, "9. Развертывание и инфраструктура", 1)
    rows = [
        ["Домен и HTTPS", "Nginx обслуживает silentx.ru и www.silentx.ru; HTTP перенаправляется на HTTPS; сертификат указан как managed by Certbot."],
        ["Статика", "client/dist отдается напрямую Nginx; /assets кешируются на 1 год как immutable."],
        ["API", "/api проксируется на Node.js-сервер http://127.0.0.1:3002."],
        ["WebSocket", "/socket.io проксируется на тот же backend с Upgrade-заголовками и долгими timeout."],
        ["База данных", "MySQL используется как основное persistent-хранилище; phpMyAdmin доступен через защищенный basic auth путь."],
        ["Сервис", "systemd-unit brenkschat.service запускает /var/www/brenkschat/server/dist/index.js."],
        ["Резервные копии", "deploy/backup содержит Node.js-скрипт mysqldump + архив исходников + checksums + retention 7 дней."],
    ]
    table(doc, "Таблица 7 — Инфраструктурные элементы", ["Элемент", "Описание"], rows, [2100, 7260])


def add_tests_and_conclusion(doc: Document) -> None:
    heading(doc, "10. Сборка, тестирование и отладка", 1)
    for item in [
        "Корневая команда npm run build последовательно собирает client и server.",
        "Клиентская сборка: tsc -b && vite build.",
        "Серверная сборка: tsc.",
        "Серверные тесты запускаются командой npm test -w server.",
        "В проекте есть тесты password.test.ts, sessions.test.ts, e2eeDevices.test.ts, store.security.test.ts.",
        "Тестами покрыты Argon2id-пароли, legacy-миграция паролей, сессии remember me, E2EE-устройства и резервные копии ключей.",
    ]:
        bullet(doc, item)

    heading(doc, "11. Итоговое описание данных", 1)
    p(
        doc,
        "Основные данные проекта делятся на пользовательские данные, коммуникационные данные, служебные "
        "настройки и защитные данные. Пользовательские данные включают профиль, аватар, почту, телефон, "
        "дату рождения и privacy-настройки. Коммуникационные данные включают чаты, участников, сообщения, "
        "медиа, реакции, ответы, статусы прочтения и закрепления. Служебные данные включают сессии, push-подписки, "
        "mute/pin-настройки и legacy-состояние. Защитные данные включают Argon2id-хэши паролей, E2EE-устройства "
        "и зашифрованные резервные копии ключей.",
    )

    heading(doc, "12. Заключение", 1)
    p(
        doc,
        "Проект «БренксЧат» является полнофункциональным веб-мессенджером с современной клиент-серверной "
        "архитектурой. В нем используются React и TypeScript для интерфейса, Node.js/Express и Socket.IO "
        "для серверной части и обмена в реальном времени, MySQL для постоянного хранения данных, а также "
        "дополнительные механизмы безопасности, push-уведомлений, E2EE, WebRTC-звонков и production-деплоя. "
        "База данных состоит из 11 таблиц, которые покрывают пользователей, чаты, сообщения, сессии, уведомления "
        "и ключи шифрования. Такая структура позволяет развивать проект дальше: выделить файловое хранилище для "
        "медиа, расширить роли пользователей, улучшить аудит безопасности и масштабировать realtime-сервер.",
    )


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_overview(doc)
    add_stack(doc)
    add_structure(doc)
    add_database(doc)
    add_features(doc)
    add_api_and_realtime(doc)
    add_security_and_deploy(doc)
    add_tests_and_conclusion(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
